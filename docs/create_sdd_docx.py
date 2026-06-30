from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "Software_Design_and_Architecture_Document.md"
TARGET = ROOT / "Software_Design_and_Architecture_Document.docx"


def add_table(document: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|") or not line.endswith("|"):
            continue
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if all(set(cell) <= {"-", ":", " "} for cell in cells):
            continue
        rows.append(cells)
    if not rows:
        return

    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=1, cols=column_count)
    table.style = "Table Grid"
    for index, value in enumerate(rows[0]):
        table.rows[0].cells[index].text = value
    for row in rows[1:]:
        cells = table.add_row().cells
        for index in range(column_count):
            cells[index].text = row[index] if index < len(row) else ""


def add_code_block(document: Document, code: str, language: str) -> None:
    title = document.add_paragraph()
    run = title.add_run(f"{language or 'code'} block")
    run.bold = True
    paragraph = document.add_paragraph()
    paragraph.style = "No Spacing"
    code_run = paragraph.add_run(code.rstrip())
    code_run.font.name = "Consolas"


def build_docx() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    document.core_properties.title = "Software Design Document and Architecture Document"
    document.core_properties.subject = "Version Manager: LangGraph + MCP Multi-Agent System"

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    table_buffer: list[str] = []
    code_buffer: list[str] = []
    in_code = False
    code_language = ""

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            add_table(document, table_buffer)
            table_buffer = []

    for line in lines:
        if line.startswith("```"):
            if in_code:
                add_code_block(document, "\n".join(code_buffer), code_language)
                code_buffer = []
                code_language = ""
                in_code = False
            else:
                flush_table()
                in_code = True
                code_language = line.strip("`").strip()
            continue

        if in_code:
            code_buffer.append(line)
            continue

        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_buffer.append(line)
            continue

        flush_table()

        stripped = line.strip()
        if not stripped:
            document.add_paragraph()
        elif stripped.startswith("# "):
            document.add_heading(stripped[2:], level=0)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:], level=2)
        elif stripped.startswith("#### "):
            document.add_heading(stripped[5:], level=3)
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped[0:2].isdigit() and ". " in stripped[:5]:
            document.add_paragraph(stripped.split(". ", 1)[1], style="List Number")
        else:
            document.add_paragraph(stripped)

    flush_table()
    if code_buffer:
        add_code_block(document, "\n".join(code_buffer), code_language)

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.save(TARGET)


if __name__ == "__main__":
    build_docx()
