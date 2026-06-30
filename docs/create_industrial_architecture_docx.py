from __future__ import annotations

from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
DIAGRAM_DIR = ROOT / "diagrams"
TARGET = ROOT / "Version_Manager_Industrial_Architecture_Diagrams.docx"

NAVY = "#17324D"
BLUE = "#2563A8"
LIGHT_BLUE = "#EAF2FB"
GREEN = "#1F7A4D"
LIGHT_GREEN = "#EAF7EF"
ORANGE = "#B45F06"
LIGHT_ORANGE = "#FFF4E5"
GRAY = "#5F6B7A"
LIGHT_GRAY = "#F3F5F7"
RED = "#A93226"
LIGHT_RED = "#FDEDEC"
BLACK = "#1F2933"
WHITE = "#FFFFFF"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


F_TITLE = font(34, True)
F_SUBTITLE = font(20, True)
F_TEXT = font(18)
F_SMALL = font(15)
F_TINY = font(13)


def draw_wrapped(draw: ImageDraw.ImageDraw, text: str, box: tuple[int, int, int, int], fill=BLACK, fnt=F_TEXT, align="center") -> None:
    x1, y1, x2, y2 = box
    max_chars = max(8, int((x2 - x1) / (fnt.size * 0.52)))
    lines: list[str] = []
    for part in text.split("\n"):
        lines.extend(wrap(part, max_chars) or [""])
    line_h = int(fnt.size * 1.25)
    total_h = line_h * len(lines)
    y = y1 + max(0, ((y2 - y1) - total_h) // 2)
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        w = bbox[2] - bbox[0]
        if align == "left":
            x = x1 + 12
        else:
            x = x1 + ((x2 - x1) - w) // 2
        draw.text((x, y), line, fill=fill, font=fnt)
        y += line_h


def box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], text: str, fill: str, outline: str, title: bool = False) -> None:
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    draw_wrapped(draw, text, xy, fill=BLACK, fnt=F_SUBTITLE if title else F_TEXT)


def band(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], title: str, fill: str, outline: str) -> None:
    draw.rounded_rectangle(xy, radius=22, fill=fill, outline=outline, width=3)
    draw.text((xy[0] + 18, xy[1] + 12), title, fill=outline, font=F_SUBTITLE)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color=GRAY, width=4, label: str | None = None) -> None:
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    size = 14
    tip = end
    p1 = (x2 - ux * size + px * size * 0.55, y2 - uy * size + py * size * 0.55)
    p2 = (x2 - ux * size - px * size * 0.55, y2 - uy * size - py * size * 0.55)
    draw.polygon([tip, p1, p2], fill=color)
    if label:
        mx, my = (x1 + x2) // 2, (y1 + y2) // 2
        label_box = (mx - 110, my - 18, mx + 110, my + 18)
        draw.rounded_rectangle(label_box, radius=8, fill=WHITE, outline="#D0D7DE")
        draw_wrapped(draw, label, label_box, fill=BLACK, fnt=F_TINY)


def base_canvas(title: str, subtitle: str = "") -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (1800, 1100), WHITE)
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1800, 88), fill=NAVY)
    draw.text((42, 22), title, fill=WHITE, font=F_TITLE)
    if subtitle:
        draw.text((42, 94), subtitle, fill=GRAY, font=F_SMALL)
    return img, draw


def save(img: Image.Image, name: str) -> Path:
    DIAGRAM_DIR.mkdir(exist_ok=True)
    path = DIAGRAM_DIR / name
    img.save(path, "PNG", quality=95)
    return path


def system_context() -> Path:
    img, d = base_canvas("System Context Architecture", "Version Manager enterprise boundaries and external integrations")
    band(d, (500, 180, 1300, 870), "Version Manager Platform", LIGHT_BLUE, BLUE)
    box(d, (565, 270, 805, 370), "Streamlit\nOperator UI", WHITE, BLUE)
    box(d, (890, 270, 1130, 370), "FastMCP\nTool Server", WHITE, BLUE)
    box(d, (565, 465, 805, 565), "LangGraph\nMulti-Agent Workflow", WHITE, BLUE)
    box(d, (890, 465, 1130, 565), "Core Domain\nServices", WHITE, BLUE)
    box(d, (565, 665, 805, 765), "SQLite Audit\nMemory", WHITE, BLUE)
    box(d, (890, 665, 1130, 765), "JSON Reports\nand Cache", WHITE, BLUE)

    box(d, (75, 260, 315, 350), "Admin\nRelease Engineer\nQA Engineer", LIGHT_GREEN, GREEN)
    box(d, (75, 465, 315, 555), "Claude Desktop\nMCP Client", LIGHT_GREEN, GREEN)
    box(d, (75, 670, 315, 760), "CLI / Scheduled\nAutomation", LIGHT_GREEN, GREEN)

    box(d, (1440, 165, 1695, 250), "OpenAI API", LIGHT_ORANGE, ORANGE)
    box(d, (1440, 295, 1695, 380), "Tavily Search", LIGHT_ORANGE, ORANGE)
    box(d, (1440, 425, 1695, 510), "NVD CVE API", LIGHT_ORANGE, ORANGE)
    box(d, (1440, 555, 1695, 640), "Enterprise Servers\nSSH / HTTP", LIGHT_ORANGE, ORANGE)
    box(d, (1440, 685, 1695, 770), "SMTP Relay", LIGHT_ORANGE, ORANGE)
    box(d, (1440, 815, 1695, 900), "Input Files\nYAML / PDF / XLSX", LIGHT_ORANGE, ORANGE)

    for s, e, lbl in [
        ((315, 305), (565, 320), "operate"),
        ((315, 510), (890, 320), "MCP"),
        ((315, 715), (890, 320), "run"),
        ((805, 320), (890, 320), "tools"),
        ((1010, 370), (685, 465), "orchestrate"),
        ((1010, 370), (1010, 465), "invoke"),
        ((685, 565), (685, 665), "audit"),
        ((1010, 565), (1010, 665), "persist"),
        ((1130, 510), (1440, 208), "LLM"),
        ((1130, 510), (1440, 338), "search"),
        ((1130, 510), (1440, 468), "CVE"),
        ((1130, 510), (1440, 598), "query"),
        ((1130, 510), (1440, 728), "notify"),
        ((1130, 510), (1440, 858), "read"),
    ]:
        arrow(d, s, e, label=lbl)
    return save(img, "01_system_context.png")


def logical_architecture() -> Path:
    img, d = base_canvas("Logical Component Architecture", "Layered design with MCP tools, agents, services, data, and integrations")
    layers = [
        (90, 150, 1710, 275, "Presentation and Clients", LIGHT_GREEN, GREEN, ["Streamlit UI", "Claude MCP Host", "CLI Runner"]),
        (90, 315, 1710, 460, "Tool and Orchestration Layer", LIGHT_BLUE, BLUE, ["FastMCP Server", "APScheduler", "LangGraph Workflow", "ReAct Agent"]),
        (90, 510, 1710, 690, "Core Domain Services", LIGHT_ORANGE, ORANGE, ["VersionFetcher", "ServerQuerier", "PDFReader", "Comparator", "VulnerabilityChecker", "ExcelReporter", "Notifier"]),
        (90, 745, 1710, 885, "Persistence and Configuration", LIGHT_GRAY, GRAY, ["config.json", "software.yml", "SQLite Memory", "JSON Reports", "JSON Cache"]),
        (90, 930, 1710, 1040, "External Integrations", LIGHT_RED, RED, ["OpenAI", "Tavily", "NVD", "SSH/HTTP Servers", "SMTP"]),
    ]
    centers = []
    for x1, y1, x2, y2, title, fill, outline, items in layers:
        band(d, (x1, y1, x2, y2), title, fill, outline)
        gap = (x2 - x1 - 70) // len(items)
        row_centers = []
        for i, item in enumerate(items):
            bx1 = x1 + 30 + i * gap
            bx2 = bx1 + gap - 22
            by1 = y1 + 58
            by2 = y2 - 18
            box(d, (bx1, by1, bx2, by2), item, WHITE, outline)
            row_centers.append(((bx1 + bx2) // 2, (by1 + by2) // 2))
        centers.append(row_centers)
    for i in range(len(centers) - 1):
        arrow(d, (900, layers[i][3]), (900, layers[i + 1][1]), label="controlled calls")
    return save(img, "02_logical_architecture.png")


def pipeline_sequence() -> Path:
    img, d = base_canvas("End-to-End Pipeline Sequence", "Run assessment from request through report and notification")
    actors = ["User / Scheduler", "MCP Server", "Current Discovery", "Latest Discovery", "Compare + Risk", "Report + Notify", "Audit / Output"]
    x_positions = [130, 390, 650, 910, 1170, 1430, 1660]
    top = 170
    for x, actor in zip(x_positions, actors):
        box(d, (x - 95, top, x + 95, top + 70), actor, LIGHT_BLUE if actor != "User / Scheduler" else LIGHT_GREEN, BLUE if actor != "User / Scheduler" else GREEN)
        d.line((x, top + 70, x, 1010), fill="#CBD5E1", width=3)
    steps = [
        (0, 1, 290, "run_full_pipeline(category)"),
        (1, 2, 390, "load software, query SSH/HTTP"),
        (2, 1, 480, "current versions or PDF fallback"),
        (1, 3, 570, "vendor/search/LLM latest lookup"),
        (3, 1, 660, "latest versions"),
        (1, 4, 750, "compare + CVE + policy"),
        (4, 1, 840, "risk findings"),
        (1, 5, 915, "Excel + HTML email"),
        (1, 6, 990, "JSON, cache, SQLite audit"),
    ]
    for src, dst, y, label in steps:
        arrow(d, (x_positions[src], y), (x_positions[dst], y), label=label)
    return save(img, "03_pipeline_sequence.png")


def data_flow() -> Path:
    img, d = base_canvas("Data Flow Diagram", "Major data stores, processing stages, and external data movement")
    process = [
        (700, 170, 1100, 260, "P1\nOrchestrate Pipeline"),
        (330, 360, 650, 450, "P2\nDiscover Current Versions"),
        (1150, 360, 1470, 450, "P3\nDiscover Latest Versions"),
        (700, 550, 1100, 640, "P4\nCompare and Assess Risk"),
        (330, 760, 650, 850, "P5\nRelease and QA Artifacts"),
        (1150, 760, 1470, 850, "P6\nNotify Stakeholders"),
    ]
    for p in process:
        box(d, p[:4], p[4], WHITE, BLUE)
    stores = [
        (75, 170, 275, 245, "D1\nconfig.json"),
        (75, 275, 275, 350, "D2\nsoftware.yml"),
        (75, 485, 275, 560, "D3\nCurrent PDF"),
        (75, 690, 275, 765, "D4\nTest Cases"),
        (1525, 170, 1735, 245, "D5\nJSON Cache"),
        (1525, 275, 1735, 350, "D6\nSQLite Memory"),
        (1525, 485, 1735, 560, "D7\nReports"),
        (1525, 690, 1735, 765, "D8\nSMTP"),
    ]
    for s in stores:
        box(d, s[:4], s[4], LIGHT_GRAY, GRAY)
    externals = [
        (330, 930, 650, 1015, "Enterprise Servers"),
        (740, 930, 1060, 1015, "Tavily / OpenAI"),
        (1150, 930, 1470, 1015, "NVD API"),
    ]
    for e in externals:
        box(d, e[:4], e[4], LIGHT_ORANGE, ORANGE)
    arrows = [
        ((275, 207), (700, 215), "settings"),
        ((275, 312), (700, 215), "inventory"),
        ((900, 260), (490, 360), "software list"),
        ((490, 450), (490, 930), "query"),
        ((275, 522), (330, 405), "fallback"),
        ((900, 260), (1310, 360), "software list"),
        ((1310, 450), (900, 550), "latest"),
        ((900, 640), (490, 760), "readiness"),
        ((275, 727), (330, 805), "test mapping"),
        ((900, 640), (1310, 760), "report data"),
        ((1310, 850), (1630, 727), "email"),
        ((1100, 595), (1525, 312), "audit"),
        ((1100, 595), (1525, 522), "reports"),
        ((1310, 360), (1630, 207), "cache"),
        ((900, 640), (1310, 930), "CVE"),
        ((1310, 360), (900, 930), "research"),
    ]
    for s, e, lbl in arrows:
        arrow(d, s, e, label=lbl)
    return save(img, "04_data_flow.png")


def deployment_architecture() -> Path:
    img, d = base_canvas("Enterprise Deployment Architecture", "Recommended target deployment boundaries")
    band(d, (80, 170, 520, 840), "User Zone", LIGHT_GREEN, GREEN)
    box(d, (160, 280, 440, 380), "Browser\nStreamlit Access", WHITE, GREEN)
    box(d, (160, 500, 440, 600), "MCP Client\nClaude Desktop", WHITE, GREEN)

    band(d, (640, 170, 1160, 840), "Application Zone", LIGHT_BLUE, BLUE)
    box(d, (745, 260, 1055, 360), "Streamlit Service\nRBAC + Dashboards", WHITE, BLUE)
    box(d, (745, 450, 1055, 550), "MCP Service\nTool Contracts", WHITE, BLUE)
    box(d, (745, 640, 1055, 740), "Worker / Scheduler\nPipeline Execution", WHITE, BLUE)

    band(d, (1280, 170, 1720, 840), "Data and Integration Zone", LIGHT_ORANGE, ORANGE)
    box(d, (1350, 245, 1650, 330), "Managed Secrets\nKey Vault", WHITE, ORANGE)
    box(d, (1350, 375, 1650, 460), "Database\nAudit + History", WHITE, ORANGE)
    box(d, (1350, 505, 1650, 590), "Object Storage\nReports + Artifacts", WHITE, ORANGE)
    box(d, (1350, 635, 1650, 720), "Redis Cache\nProvider Results", WHITE, ORANGE)
    box(d, (1350, 765, 1650, 835), "External APIs\nOpenAI / Tavily / NVD / SMTP", WHITE, ORANGE)

    for s, e, lbl in [
        ((440, 330), (745, 310), "HTTPS"),
        ((440, 550), (745, 500), "MCP"),
        ((900, 360), (900, 450), "invoke"),
        ((900, 550), (900, 640), "jobs"),
        ((1055, 500), (1350, 287), "secrets"),
        ((1055, 500), (1350, 417), "audit"),
        ((1055, 690), (1350, 547), "artifacts"),
        ((1055, 690), (1350, 677), "cache"),
        ((1055, 690), (1350, 800), "integrate"),
    ]:
        arrow(d, s, e, label=lbl)
    return save(img, "05_deployment_architecture.png")


def set_styles(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = "Segoe UI"
    styles["Normal"].font.size = Pt(10.5)
    for name, size, color in [
        ("Title", 24, NAVY),
        ("Heading 1", 17, NAVY),
        ("Heading 2", 13, BLUE),
    ]:
        style = styles[name]
        style.font.name = "Segoe UI"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color.replace("#", ""))


def add_caption(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(95, 107, 122)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Shading Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def build_doc() -> None:
    diagrams = [
        ("Figure 1. System Context Architecture", system_context()),
        ("Figure 2. Logical Component Architecture", logical_architecture()),
        ("Figure 3. End-to-End Pipeline Sequence", pipeline_sequence()),
        ("Figure 4. Data Flow Diagram", data_flow()),
        ("Figure 5. Enterprise Deployment Architecture", deployment_architecture()),
    ]

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    set_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Version Manager Architecture Diagram Pack")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor.from_string(NAVY.replace("#", ""))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("LangGraph + MCP Multi-Agent System | Enterprise Architecture Views")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(95, 107, 122)

    doc.add_heading("Architecture Overview", level=1)
    doc.add_paragraph(
        "This document provides industrial-style architecture diagrams for the Version Manager platform. "
        "The diagrams are embedded as high-resolution images so they can be reviewed directly in Word, "
        "shared with stakeholders, and reused in enterprise design reviews without Mermaid rendering."
    )

    add_table(
        doc,
        ["View", "Purpose", "Primary Audience"],
        [
            ["System Context", "Shows users, platform boundary, and external integrations.", "Executives, architects, security"],
            ["Logical Architecture", "Shows layered components and module responsibilities.", "Architects, developers"],
            ["Pipeline Sequence", "Shows execution order for the full assessment flow.", "Developers, operators"],
            ["Data Flow Diagram", "Shows data movement across processes and stores.", "Security, audit, architecture"],
            ["Deployment Architecture", "Shows recommended enterprise deployment zones.", "Infrastructure, DevOps, security"],
        ],
    )

    for caption, path in diagrams:
        doc.add_page_break()
        doc.add_heading(caption.split(". ", 1)[1], level=1)
        doc.add_picture(str(path), width=Inches(10.6))
        add_caption(doc, caption)

    doc.add_page_break()
    doc.add_heading("Component Responsibility Summary", level=1)
    add_table(
        doc,
        ["Component", "Enterprise Responsibility"],
        [
            ["Streamlit UI", "Operator dashboard, report visualization, role-aware workflows."],
            ["FastMCP Server", "Stable tool interface, deterministic pipeline coordination, scheduler lifecycle."],
            ["LangGraph Workflow", "Least-privilege multi-agent orchestration with explicit state transitions."],
            ["Core Services", "Version discovery, comparison, security assessment, reporting, notification."],
            ["SQLite / Reports / Cache", "Audit memory, generated artifacts, provider result reuse."],
            ["External Integrations", "OpenAI, Tavily, NVD, enterprise servers, and SMTP relay."],
        ],
    )

    doc.save(TARGET)


if __name__ == "__main__":
    build_doc()
